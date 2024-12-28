import frappe
from frappe.model.document import Document
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from datetime import datetime, timedelta
import os
from frappe.utils import get_files_path
from urllib.parse import urljoin
import time
import random
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException

def is_valid_subpath(base_url, link):
    if not link or link.startswith(('#', 'javascript:', 'mailto:', 'tel:')):
        return False
    try:
        full_url = urljoin(base_url, link)
        return full_url.startswith(base_url)
    except:
        return False

class WebCrawler(Document):
    def validate(self):
        if not self.website_url.startswith(('http://', 'https://')):
            frappe.throw('Please enter a valid URL starting with http:// or https://')

    def before_save(self):
        if not self.status:
            self.status = "Pending"

    def setup_selenium(self):
        try:
            print("=== Debug: Starting setup_selenium ===")
            chrome_options = Options()
            chrome_options.add_argument('--headless')
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            print("=== Debug: Basic chrome options set ===")
            
            # Add these memory and process related options
            chrome_options.add_argument('--disable-gpu')
            chrome_options.add_argument('--disable-software-rasterizer')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--no-zygote')
            chrome_options.add_argument('--single-process')
            chrome_options.add_argument('--disable-infobars')
            chrome_options.add_argument('--disable-notifications')
            print("=== Debug: Performance options set ===")
            
            # Add these new options for better timeout handling
            chrome_options.add_argument('--dns-prefetch-disable')
            chrome_options.add_argument('--disable-extensions')
            chrome_options.page_load_strategy = 'eager'
            print("=== Debug: Additional chrome options set ===")
            
            # Try to find Chrome/Chromium binary in different possible locations
            browser_paths = [
                '/usr/bin/chromium',
                '/usr/bin/google-chrome',
                '/usr/bin/chromium-browser',
                '/snap/bin/chromium'
            ]
            
            print("=== Debug: Searching for browser binary ===")
            binary_found = False
            for path in browser_paths:
                print(f"=== Debug: Checking path: {path} ===")
                if os.path.exists(path):
                    chrome_options.binary_location = path
                    binary_found = True
                    print(f"=== Debug: Browser found at: {path} ===")
                    break
                
            if not binary_found:
                print("=== Debug: No browser binary found ===")
                frappe.throw("No Chrome or Chromium browser found. Please install either Chrome or Chromium.")
            
            print("=== Debug: Setting additional chrome arguments ===")
            chrome_options.add_argument('--disable-gpu')
            chrome_options.add_argument('--window-size=1920x1080')
            chrome_options.add_argument(f'user-agent={self.get_random_user_agent()}')
            chrome_options.add_argument('--disable-blink-features=AutomationControlled')
            chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
            chrome_options.add_experimental_option('useAutomationExtension', False)
            
            print("=== Debug: Starting ChromeDriverManager installation ===")
            service = Service(ChromeDriverManager().install())
            print("=== Debug: ChromeDriverManager installation complete ===")
            
            print("=== Debug: Creating Chrome driver ===")
            try:
                driver = webdriver.Chrome(
                    service=service, 
                    options=chrome_options
                )
                print("=== Debug: Chrome driver created successfully ===")
            except Exception as e:
                print(f"=== Debug: Error creating Chrome driver: {str(e)} ===")
                # Kill any hanging chromium processes
                os.system('pkill chromium')
                os.system('pkill chrome')
                raise e
            
            driver.set_page_load_timeout(300)
            driver.set_script_timeout(300)
            
            print("=== Debug: setup_selenium completed successfully ===")
            return driver
            
        except Exception as e:
            print(f"=== Debug: Final error in setup_selenium: {str(e)} ===")
            # Cleanup any hanging processes
            os.system('pkill chromium')
            os.system('pkill chrome')
            raise e

    def get_random_user_agent(self):
        user_agents = [
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:89.0) Gecko/20100101 Firefox/89.0',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/14.1.1 Safari/605.1.15',
        ]
        return random.choice(user_agents)

    @frappe.whitelist()
    def start_crawling(self):
        try:
            print("=== Debug: start_crawling method called ===")
            self.status = "Processing"
            self.save()
            
            print("=== Debug: Setting up Selenium ===")
            driver = self.setup_selenium()
            print("=== Debug: Selenium setup complete ===")
            
            try:
                start_time = datetime.now()
                visited = set()
                seen_content = set()
                max_depth = int(self.max_depth) if self.max_depth else 2
                timeout = float(self.time_limit) if self.time_limit else 300
                
                doc = DocxDocument()
                doc.add_heading(f'Crawled Content from {self.website_url}', 0)
                doc.add_paragraph(f'Crawled on: {start_time.strftime("%Y-%m-%d %H:%M:%S")}')

                print(f"=== Debug: About to call crawl method with URL: {self.website_url} ===")
                self.crawl(
                    self.website_url,
                    depth=0,
                    max_depth=max_depth,
                    visited=visited,
                    document=doc,
                    start_time=start_time,
                    seen_content=seen_content,
                    timeout_minutes=timeout/60,
                    driver=driver
                )
                print("=== Debug: Crawl method completed ===")

            finally:
                driver.quit()
                print("=== Debug: Driver quit successfully ===")

            # Save the document
            file_name = f'crawled_content_{frappe.utils.now_datetime().strftime("%Y%m%d_%H%M%S")}.docx'
            file_path = os.path.join(get_files_path(), file_name)
            doc.save(file_path)

            self.output_file = f'/files/{file_name}'
            self.status = "Completed"
            self.save()

            frappe.msgprint("Crawling completed successfully!")

        except Exception as e:
            frappe.logger().error(f"Error in start_crawling: {str(e)}")
            self.status = "Failed"
            self.save()
            frappe.log_error(f"Web Crawler Error: {str(e)}")
            frappe.throw(f"Error during crawling: {str(e)}")

    def crawl(self, url, depth, max_depth, visited, document, start_time, seen_content, timeout_minutes, driver):
        print(f"Crawling ({depth}/{max_depth}): {url}")
        if datetime.now() - start_time > timedelta(minutes=timeout_minutes):
            frappe.logger().warning(f"Timeout reached ({timeout_minutes} minutes). Stopping crawl process.")
            return

        if depth > max_depth:
            frappe.logger().debug(f"Max depth {max_depth} reached for {url}")
            return
            
        if url in visited:
            frappe.logger().debug(f"URL already visited: {url}")
            return

        frappe.logger().info(f"Crawling ({depth}/{max_depth}): {url}")
        frappe.publish_realtime('web_crawler_progress', {
            'url': url,
            'depth': depth,
            'max_depth': max_depth,
            'visited_count': len(visited)
        })
        
        visited.add(url)

        try:
            delay = random.uniform(2, 5) * (depth + 1)
            frappe.logger().debug(f"Waiting {delay:.2f} seconds before crawling {url}")
            time.sleep(delay)

            try:
                driver.get(url)
                frappe.logger().debug(f"Page loaded: {url}")
                WebDriverWait(driver, 30).until(
                    EC.presence_of_all_elements_located((By.TAG_NAME, "body"))
                )
            except TimeoutException:
                frappe.logger().error(f"Timeout while loading {url}, skipping...")
                document.add_paragraph(f"Timeout while loading {url}, skipped")
                driver.execute_script('window.stop();')
                return
            
            soup = BeautifulSoup(driver.page_source, "html.parser")
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            content = soup.get_text(separator=' ', strip=True)
            words = content.split()
            if len(words) > 100:
                content = ' '.join(words[:100]) + '...'
                frappe.logger().debug(f"Content truncated to 100 words for {url}")

            content_preview = ' '.join(words[:20])
            if content_preview not in seen_content:
                seen_content.add(content_preview)
                document.add_heading(f'Content from: {url}', level=1)
                document.add_paragraph(content)
                document.add_paragraph('-------------------')
                frappe.logger().info(f"Saved unique content from {url}")
            else:
                frappe.logger().debug(f"Skipping duplicate content from {url}")

            links = [a.get_attribute('href') for a in driver.find_elements(By.TAG_NAME, "a")]
            subpaths = [urljoin(url, link) for link in links if is_valid_subpath(url, link)]
            
            frappe.logger().debug(f"Found {len(subpaths)} valid links on {url}")

            for subpath in subpaths:
                if datetime.now() - start_time > timedelta(minutes=timeout_minutes):
                    frappe.logger().warning(f"Timeout reached ({timeout_minutes} minutes). Stopping crawl process.")
                    return
                self.crawl(subpath, depth + 1, max_depth, visited, document, start_time, seen_content, timeout_minutes, driver)

        except (TimeoutException, WebDriverException) as e:
            error_msg = f"Error crawling {url}: {str(e)}"
            frappe.logger().error(error_msg)
            document.add_paragraph(error_msg)
        except Exception as e:
            error_msg = f"Unexpected error crawling {url}: {str(e)}"
            frappe.logger().error(error_msg)
            document.add_paragraph(error_msg)