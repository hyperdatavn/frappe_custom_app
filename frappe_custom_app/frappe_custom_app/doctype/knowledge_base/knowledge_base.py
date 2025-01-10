import frappe
import os
import json
import numpy as np
from langchain.embeddings import OpenAIEmbeddings
import faiss
from docx import Document as DocxDocument
from frappe.model.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

CHUNK_SIZE = 500
FAISS_INDEX_PATH = "private/files/faiss_index.bin"
FAISS_METADATA_PATH = "private/files/faiss_metadata.json"

class KnowledgeBase(Document):
    """Knowledge Base Document Class"""
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.api_key = frappe.db.get_single_value("Hyperdata App Settings", "api_key")

    @frappe.whitelist()
    def store_content_as_vector(self):
        """
        Process the uploaded .docx file, extract text, and store it as a vector embedding.
        """
        file_url = self.document_file

        if not file_url:
            frappe.throw("Please upload a document file before preparing data.")

        file_path = frappe.get_site_path("private", file_url.replace("/private/files/", "files/"))
        if os.path.exists(file_path):
            print("File exists at:", file_path)
        else:
            print("File not found!")
               
        # Generate vector embeddings
        if not self.api_key:
            frappe.throw("OpenAI API key is missing. Please configure it in Hyperdata App Settings.")
        # Split document into chunks
        if self.status == "Draft":
            chunks = self._chunk_document(file_path)
            self.status = "Chunked"
            self.save()
        else:
            chunks = self._chunk_document(file_path)

        # Generate vectors for each chunk
        for chunk_idx, chunk in enumerate(chunks):
            document_vector = self._generate_vector_from_text(chunk['text'], self.api_key)
            
            # Enhanced metadata per chunk
            metadata = {
                "file_name": os.path.basename(file_path),
                "source": self.source,
                "chunk_index": chunk_idx,
                "section_title": chunk.get('section_title', ''),
                "description": chunk['text'][:100],
                "content_length": len(chunk['text']),
                "update_date": frappe.utils.now()
            }
            
            self._store_vector_in_faiss(document_vector, metadata)

            # Update the status of corresponding Knowledge Base Chunk
            chunk_doc = frappe.get_all(
                "Knowledge Base Chunk",
                filters={
                    "knowledge_base": self.name,
                    "chunk_index": chunk_idx
                },
                pluck="name"
            )
            if chunk_doc:
                frappe.db.set_value("Knowledge Base Chunk", chunk_doc[0], "status", "Processed")

        # Update document status
        self.processed_data = f"Vector stored with metadata: {metadata}"
        self.status = "Completed"
        self.save()

    def _extract_text_from_document(self, file_path):
        try:
            document = DocxDocument(file_path)
            text = []
            for paragraph in document.paragraphs:
                if paragraph.style.name.startswith("Heading"):  # Ignore heading
                    continue
                if paragraph.text.strip():
                    text.append(paragraph.text.strip())
            return self.filter_text("\n".join(text))
        except Exception as e:
            frappe.throw(f"Error reading the document file: {str(e)}")

    def filter_text(self, text):
        """
        Filter out irrelevant or low-quality text such as headings or boilerplate content.
        """
        filtered_lines = []
        for line in text.split("\n"):
            if len(line.strip()) > 10 and any(c.isalnum() for c in line):
                filtered_lines.append(line.strip())
        return "\n".join(filtered_lines)        

    def _generate_vector_from_text(self, text, api_key):
        """
        Generate a vector embedding for the given text using OpenAI.
        """
        try:
            embeddings = OpenAIEmbeddings(openai_api_key=api_key)
            return embeddings.embed_query(text)
        except Exception as e:
            frappe.throw(f"Failed to generate embeddings: {str(e)}")

    def _store_vector_in_faiss(self, document_vector, metadata):
        """
        Store vector embeddings and metadata into FAISS index.
        """
        try:
            faiss_index_path = frappe.get_site_path("private", "files", "faiss_index.bin")
            metadata_path = frappe.get_site_path("private", "files", "faiss_metadata.json")

            # Load or initialize FAISS index
            if os.path.exists(faiss_index_path):
                index = faiss.read_index(faiss_index_path)
                print("Number of vectors in FAISS Index:", index.ntotal)

            else:
                index = faiss.IndexFlatL2(len(document_vector))

            # Load or initialize metadata
            if os.path.exists(metadata_path):
                with open(metadata_path, "r", encoding="utf-8") as f:
                    metadata_list = json.load(f)
            else:
                metadata_list = []

            # Add vector and metadata
            if index.is_trained:
                index.add(np.array([document_vector], dtype="float32"))
                metadata_list.append(metadata)
            else:
                frappe.throw("FAISS Index is not trained. Failed to add vector.")


            # Save updated index and metadata
            faiss.write_index(index, faiss_index_path)
            with open(metadata_path, "w", encoding="utf-8") as f:
                json.dump(metadata_list, f, ensure_ascii=False, indent=4)

        except Exception as e:
            frappe.throw(f"Failed to store vector in FAISS: {str(e)}")

    def _chunk_document(self, file_path, chunk_size=500, chunk_overlap=50):
        """Split document into chunks using LangChain's TextSplitter"""

        # Read the entire document text
        try:
            document = DocxDocument(file_path)
            full_text = []
            current_section = None

            for paragraph in document.paragraphs:
                # Section headings
                if paragraph.style.name.startswith("Heading"):
                    current_section = paragraph.text
                    continue
                text = paragraph.text.strip()
                if text:
                    full_text.append(text)

            text = "\n".join(full_text)
        except Exception as e:
            frappe.throw(f"Error reading the document file: {str(e)}")

        # Use LangChain's TextSplitter
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )

        # Split text into chunks
        chunks = text_splitter.create_documents([text])

        # Prepare chunks with metadata
        chunk_list = []
        for idx, chunk in enumerate(chunks):
            chunk_list.append({
                'text': chunk.page_content,
                'section_title': current_section
            })

        return chunk_list

    @frappe.whitelist()
    def create_content_chunks(self):
        """
        Process the uploaded .docx file and create chunk documents for review
        """
        file_url = self.document_file
        if not file_url:
            frappe.throw("Please upload a document file before preparing data.")

        file_path = frappe.get_site_path("private", file_url.replace("/private/files/", "files/"))
        
        # Split document into chunks
        chunks = self._chunk_document(file_path)
        
        # Create chunk documents
        for chunk_idx, chunk in enumerate(chunks):
            chunk_doc = frappe.get_doc({
                "doctype": "Knowledge Base Chunk",
                "knowledge_base": self.name,
                "chunk_index": chunk_idx,
                "section_title": chunk.get('section_title', ''),
                "content": chunk['text'],
                "status": "Draft"
            })
            chunk_doc.insert()

        self.status = "Chunked"
        self.save()

    @frappe.whitelist()
    def process_reviewed_chunks(self):
        """
        Process all reviewed chunks and store them as vectors
        """
        if not self.api_key:
            frappe.throw("OpenAI API key is missing. Please configure it in Hyperdata App Settings.")

        # Get all reviewed chunks
        chunks = frappe.get_all(
            "Knowledge Base Chunk",
            filters={
                "knowledge_base": self.name,
                "status": "Reviewed"
            },
            fields=["name", "content", "section_title", "chunk_index"],
            order_by="chunk_index"
        )

        for chunk in chunks:
            document_vector = self._generate_vector_from_text(chunk.content, self.api_key)
            
            metadata = {
                "chunk_doc": chunk.name,
                "source": self.source,
                "chunk_index": chunk.chunk_index,
                "section_title": chunk.section_title,
                "description": chunk.content[:100],
                "content_length": len(chunk.content),
                "update_date": frappe.utils.now()
            }
            
            self._store_vector_in_faiss(document_vector, metadata)
            
            # Update chunk status
            frappe.db.set_value("Knowledge Base Chunk", chunk.name, "status", "Processed")

        self.status = "Completed"
        self.save()

@frappe.whitelist()
def store_content_as_vector(docname):
    doc = frappe.get_doc("Knowledge Base", docname)
    return doc.store_content_as_vector()

@frappe.whitelist()
def upload_file_to_knowledge_base(source):
    """
    API endpoint to upload a file to the Knowledge Base.
    :param source: The source of the document.
    """
    from frappe.utils.file_manager import save_file

    # Check if the file is present in the request
    if 'file' not in frappe.request.files:
        frappe.throw("No file found in the request")

    # Get the uploaded file from the request
    uploaded_file = frappe.request.files['file']

    # Read the file content
    file_content = uploaded_file.stream.read()  # Use .read() if 'stream' is not available

    # Create a new Knowledge Base document
    kb_doc = frappe.new_doc("Knowledge Base")
    kb_doc.source = source
    kb_doc.status = "Draft"
    kb_doc.insert()

    # Save the uploaded file and attach it to the Knowledge Base document
    saved_file = save_file(
        uploaded_file.filename,
        file_content,
        "Knowledge Base",
        kb_doc.name,
        is_private=1
    )

    # Update the document_file field with the URL of the saved file
    kb_doc.document_file = saved_file.file_url
    kb_doc.save()

    # Try to create chunks
    try:
        kb_doc.create_content_chunks()
    except Exception as e:
        # Log the error but don't throw it to allow file upload to complete
        frappe.log_error(f"Chunk creation failed for {kb_doc.name}: {str(e)}", 
                        "Knowledge Base Chunk Creation")
        return {
            "message": "File uploaded successfully but chunk creation failed",
            "knowledge_base": kb_doc.name
        }


    return {
        "message": "File uploaded and chunked successfully",
        "knowledge_base": kb_doc.name
    }
